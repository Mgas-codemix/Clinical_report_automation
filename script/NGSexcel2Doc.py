# Written in python 3.10.12
import pandas as pd
import argparse
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def create_parser():
    """Create argument parser for input and output files."""
    parser = argparse.ArgumentParser(description="Generate clinical reports from NGS Excel data.")
    parser.add_argument("-inFile", type=str, help="Path to input NGS Excel file", required=True)
    parser.add_argument("-template", type=str, help="Path to Word template file", required=True)
    parser.add_argument("-outdir", type=str, help="Path to output directory for reports", required=True) 
    return parser

def generate_text(gene, HGVScons, vaf, varID):
    """Generate formatted text for gene mutation details."""
    return f"\tGene: {gene}\n\tVariant: {HGVScons}\n\tVAF: {round(vaf*100, 2)}%\n\tVariant ID: {varID}\n\n"

def move_table_after(table, paragraph):
    """Move the table after a given paragraph in the document."""
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

def create_table(document, array, idx, after_line=False,  bold_header=False, style='TableGrid', font_size=12, autofit=False, allow_autofit=False, align_vertical=False):
    """Generates a NxN table after a specified line occurence or at the end of a document."""
    r, c = (len(array), len(array[0]))
    table = document.add_table(rows=r, cols=c)
    table.style.style_id = style
    table.autofit = autofit
    table.allow_autofit = allow_autofit

    if after_line: move_table_after(table, document.paragraphs[idx+1])

    for i in range(r):
        for j in range(c):
            if i == 0 and bold_header:
                table.rows[i].cells[j].paragraphs[0].add_run(str(array[i][j])).bold=True
            else:
                table.rows[i].cells[j].text = str(array[i][j])

    # Set font size
    for row in table.rows:
        for cell in row.cells:
            #if align_vertical: cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if align_vertical: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

def populate_report(document, dataframe, record):
    """Populate the report for each Record with PHI and Somatic Variant data."""
    for idx, paragraph in enumerate(document.paragraphs):
        if 'Clinical Report' in paragraph.text:
            # Insert patient health information
            PHI = dataframe.loc[dataframe['SampleID'] == record, ['RecordNumber', 'SampleID', 'SampleDate', 'RunDate', 'SampleType']].drop_duplicates()
            PHI.columns = ['Record Number', 'Sample ID', 'Sample Date', 'Run Date', 'Sample Type']
            PHI = pd.concat([PHI.columns.to_frame().T, PHI], sort=False)
            create_table(document, PHI.T.to_numpy(), idx, True, False, 'TableGrid', 10, True, True, False)

        elif 'Somatic Variants' in paragraph.text:
            # Insert somatic variants table
            variant_data = dataframe[dataframe['SampleID'] == record][['Chrom', 'Gene', 'Variant-ID', 'VAF', 'HGVS Consequence']]
            variant_data['VAF'] = variant_data['VAF'].multiply(100).round(2)
            variant_data = pd.concat([variant_data.columns.to_frame().T, variant_data], sort=False)
            create_table(document, variant_data.to_numpy(), idx, False, True, 'TableGrid', 8, True, True, True)
            #break  # After inserting the variant data, stop processing this document

def main():
    args = create_parser().parse_args()

    # Ensure output directory exists
    os.makedirs(args.outdir, exist_ok=True)

    # Load data from Excel
    NGSdata = pd.read_excel(args.inFile)

    # Prepare cases for each unique SampleID
    uniqueIDS = NGSdata['SampleID'].unique()
    Cases = {sample_id: Document(args.template) for sample_id in uniqueIDS}

    # Generate reports
    for sample_id, document in Cases.items():
        populate_report(document, NGSdata, sample_id)
        document.save(f"{args.outdir}/{sample_id}.docx")

if __name__ == '__main__':
    main()