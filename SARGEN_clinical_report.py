
#!/usr/bin/env python3

import pandas as pd
import argparse
import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

def create_parser():
    parser = argparse.ArgumentParser(description="Generate reports from MultiQC HTML data and Excel files")
    parser.add_argument("-inFile", dest="inFile", type=str, nargs=1, help="Path to input MultiQC HTML file", required=True)
    parser.add_argument("-template", dest="template", type=str, nargs=1, help="Path to the Word template file", required=True)
    parser.add_argument("-excelfile", dest="excelfile", type=str, nargs=1, help="Path to the Excel file with additional tables", required=True)
    parser.add_argument("-outdir", dest="output", type=str, nargs=1, help="Path to output directory where reports will be saved", required=True)
    args = parser.parse_args()
    return args

def parse_multiqc_report(multiqc_file):
    """
    Parses the MultiQC HTML report to extract specific data (e.g., coverage, quality metrics).
    :param multiqc_file: Path to the MultiQC HTML report
    :return: Dictionary containing extracted values
    """
    with open(multiqc_file, 'r') as f:
        soup = BeautifulSoup(f, 'lxml')

    # Define specific section IDs or classes to extract data
    data = {}

    # Example: Extract coverage metrics from a specific section
    coverage_section = soup.find('div', {'id': 'coverage-section'})  # Replace with actual section identifier
    if coverage_section:
        coverage_value = coverage_section.find('span', {'class': 'metric-value'}).text.strip()  # Modify accordingly
        data['coverage'] = coverage_value

    # Example: Extract quality metrics
    quality_section = soup.find('div', {'id': 'quality-section'})  # Replace with actual section identifier
    if quality_section:
        quality_value = quality_section.find('span', {'class': 'metric-value'}).text.strip()
        data['quality'] = quality_value

    return data

def extract_tables_from_excel(excelfile_path):
    """
    Extracts tables from the given Excel file.
    :param excelfile_path: Path to the Excel file
    :return: Dictionary containing extracted tables (as Pandas DataFrames)
    """
    # Read Excel file into a dictionary of dataframes
    excel_data = pd.read_excel(excelfile_path, sheet_name=None)  # Load all sheets as a dict of DataFrames
    return excel_data

def insert_table_from_dataframe(document, df):
    """
    Inserts a Pandas DataFrame as a table into the Word document.
    :param document: Word document object
    :param df: DataFrame to insert as a table
    """
    table = document.add_table(rows=df.shape[0]+1, cols=df.shape[1])

    # Insert column headers
    for j, col_name in enumerate(df.columns):
        table.cell(0, j).text = str(col_name)

    # Insert rows from the dataframe
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i+1, j).text = str(df.iloc[i, j])

def replace_placeholders(document, placeholders):
    """
    Replaces placeholders in the Word document with actual data.
    :param document: Word document object
    :param placeholders: Dictionary of placeholder keys and their values to replace
    """
    for paragraph in document.paragraphs:
        for key, value in placeholders.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

def generate_reports_from_multiqc_and_excel(multiqc_data, excel_tables, template_path, output_dir, sample_id):
    """
    Generates a report based on data extracted from MultiQC and Excel file.
    :param multiqc_data: Dictionary containing MultiQC-extracted data (e.g., coverage, quality metrics, etc.)
    :param excel_tables: Dictionary of tables extracted from the Excel file (as DataFrames)
    :param template_path: Path to the Word template file
    :param output_dir: Directory where the generated reports will be saved
    :param sample_id: Sample identifier (for report file naming)
    """
    # Load the Word template
    document = Document(template_path)

    # Define placeholders and map them to corresponding extracted values
    placeholders = {
        '<<SampleID>>': sample_id,
        '<<CoverageValue>>': multiqc_data.get('coverage', 'N/A'),  # Replace placeholder with coverage value
        '<<QualityValue>>': multiqc_data.get('quality', 'N/A'),  # Replace placeholder with quality value
    }

    # Replace placeholders in the document
    replace_placeholders(document, placeholders)

    # Insert tables from the Excel file
    for sheet_name, df in excel_tables.items():
        document.add_paragraph(f"Table from {sheet_name} sheet:")
        insert_table_from_dataframe(document, df)

    # Save the report
    output_path = os.path.join(output_dir, f'report_{sample_id}.docx')
    document.save(output_path)
    print(f'Report saved at {output_path}')

def main():
    # Parse arguments
    myargs = create_parser()
    inputFile = myargs.inFile[0]
    templatePath = myargs.template[0]
    outputDir = myargs.output[0]
    excelFile = myargs.excelfile[0]

    # Parse the MultiQC report to extract relevant data
    multiqc_data = parse_multiqc_report(inputFile)

    # Extract tables from the Excel file
    excel_tables = extract_tables_from_excel(excelFile)

    # Sample ID (assuming this is in the MultiQC report or input file; you may need to adapt this)
    sample_id = 'Sample1'  # Replace this with the actual way to retrieve sample ID

    # Generate the report using the template, extracted MultiQC data, and Excel tables
    generate_reports_from_multiqc_and_excel(multiqc_data, excel_tables, templatePath, outputDir, sample_id)

if __name__ == '__main__':
    main()
